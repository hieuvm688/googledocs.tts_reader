"""Adapter đọc nội dung từ file .docx cục bộ."""
from pathlib import Path
from src.domain.entities.document import Document
from src.domain.exceptions.base import DocumentFetchError
from src.domain.ports.document_source_port import IDocumentSource

class DocxFileAdapter(IDocumentSource):
    """Đọc tài liệu Word (.docx) hoặc file text (.txt) trên máy."""

    def can_handle(self, source_identifier: str) -> bool:
        identifier = source_identifier.strip().lower()
        return identifier.endswith(".docx") or identifier.endswith(".txt")

    def fetch_document(self, source_identifier: str) -> Document:
        file_path = Path(source_identifier).resolve()
        if not file_path.exists():
            raise DocumentFetchError(
                "Tệp tài liệu không tồn tại trên hệ thống.",
                details=str(file_path)
            )

        content_parts = []
        title = file_path.stem

        if file_path.suffix.lower() == ".docx":
            try:
                import docx
                doc = docx.Document(str(file_path))
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        content_parts.append(text)
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_texts:
                            content_parts.append(" | ".join(row_texts))
            except Exception as exc:
                raise DocumentFetchError(
                    "Không thể đọc nội dung tệp .docx.",
                    details=str(exc)
                ) from exc
        else:
            # File text thông thường
            try:
                content_parts.append(file_path.read_text(encoding="utf-8"))
            except Exception as exc:
                raise DocumentFetchError(
                    "Không thể đọc nội dung tệp text.",
                    details=str(exc)
                ) from exc

        content = "\n\n".join(content_parts).strip()
        return Document(
            id=file_path.name,
            title=title,
            content=content,
            source_type="docx_file",
            metadata={"file_path": str(file_path)}
        )
